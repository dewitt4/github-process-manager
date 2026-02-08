"""
Word Document Generator for Process Analysis Reports.
Generates professionally formatted .docx files with structured sections.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import re
from logger import logger


def parse_analysis_sections(analysis_text):
    """
    Parse the chatbot response into structured sections.
    
    Sections:
    1. Control Objective (or Model Overview, Pipeline Overview, etc.)
    2. Risks Addressed (or Data Pipeline, Build Steps, etc.)
    3. Testing Procedures
    4. Test Results and Findings
    5. Conclusion and Recommendation
    
    Args:
        analysis_text (str): Full chatbot response text
        
    Returns:
        dict: Dictionary with section titles as keys and content as values
    """
    sections = {
        'Control Objective': '',
        'Risks Addressed': '',
        'Testing Procedures': '',
        'Test Results and Findings': '',
        'Conclusion and Recommendation': ''
    }
    
    # Try to parse using section markers
    patterns = [
        (r'1\.\s*Control Objective[:\s]*(.*?)(?=2\.\s*Risks|$)', 'Control Objective'),
        (r'2\.\s*Risks Addressed[:\s]*(.*?)(?=3\.\s*Testing|$)', 'Risks Addressed'),
        (r'3\.\s*Testing Procedures[:\s]*(.*?)(?=4\.\s*Test Results|4\.\s*Results|$)', 'Testing Procedures'),
        (r'4\.\s*(?:Test Results and Findings|Results)[:\s]*(.*?)(?=5\.\s*Conclusion|$)', 'Test Results and Findings'),
        (r'5\.\s*(?:Conclusion and Recommendation|Conclusion)[:\s]*(.*?)$', 'Conclusion and Recommendation')
    ]
    
    for pattern, section_name in patterns:
        match = re.search(pattern, analysis_text, re.DOTALL | re.IGNORECASE)
        if match:
            sections[section_name] = match.group(1).strip()
    
    # If no sections were parsed, use the full text as a general control analysis
    if not any(sections.values()):
        sections['Control Objective'] = analysis_text
    
    return sections


def create_process_document(analysis_text, process_name='Process Analysis', metadata=None):
    """
    Generate Process Analysis Word document.
    
    Creates a professionally formatted Word document with:
    - Clean header with project branding
    - 5 structured sections
    - Proper styling (Calibri, headings, spacing)
    - Page footer with page numbers
    
    Args:
        analysis_text (str): The chatbot's analysis response
        process_name (str): Name of the process being analyzed
        metadata (dict): Optional metadata (timestamp, query, user)
        
    Returns:
        str: Filename of the generated document
    """
    try:
        # Create generated_reports directory if it doesn't exist
        os.makedirs('generated_reports', exist_ok=True)
        
        # Parse sections from the analysis text
        sections = parse_sox_sections(analysis_text)
        
        # Create new document
        doc = Document()
        
        # Set default font to Calibri
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add header with project branding
        header_section = doc.sections[0]
        header = header_section.header
        header_para = header.paragraphs[0]
        header_para.text = 'GitHub Process Manager | Process Documentation'
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.runs[0]
        header_run.font.size = Pt(10)
        header_run.font.color.rgb = RGBColor(74, 144, 226)  # Professional blue (#4A90E2)
        
        # Add title
        title = doc.add_heading('Process Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(74, 144, 226)  # Professional blue (#4A90E2)
        title_run.font.size = Pt(18)
        
        # Add process name
        process_heading = doc.add_heading(process_name, level=2)
        control_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add metadata section
        if metadata:
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_table.cell(0, 0).text = 'Generated Date:'
            info_table.cell(0, 1).text = metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            info_table.cell(1, 0).text = 'Analysis Query:'
            info_table.cell(1, 1).text = metadata.get('query', 'N/A')
            
            info_table.cell(2, 0).text = 'Report Type:'
            info_table.cell(2, 1).text = 'Process Analysis Documentation'
            
            doc.add_paragraph()  # Spacing
        
        # Add horizontal line
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Add the 5 sections
        section_titles = [
            ('1. Control Objective', 'Control Objective'),
            ('2. Risks Addressed', 'Risks Addressed'),
            ('3. Testing Procedures', 'Testing Procedures'),
            ('4. Test Results and Findings', 'Test Results and Findings'),
            ('5. Conclusion and Recommendation', 'Conclusion and Recommendation')
        ]
        
        for title_text, section_key in section_titles:
            # Add section heading
            section_heading = doc.add_heading(title_text, level=1)
            section_run = section_heading.runs[0]
            section_run.font.color.rgb = RGBColor(74, 144, 226)  # Professional blue (#4A90E2)
            section_run.font.size = Pt(14)
            
            # Add section content
            content = sections.get(section_key, '').strip()
            if content:
                # Check if content has bullet points or numbered lists
                if '\n-' in content or '\n•' in content or re.search(r'\n\d+\.', content):
                    # Split into lines and add as list items
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            if line.startswith('-') or line.startswith('•'):
                                doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            elif re.match(r'^\d+\.', line):
                                doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
                            else:
                                doc.add_paragraph(line)
                else:
                    doc.add_paragraph(content)
            else:
                doc.add_paragraph('[No content provided for this section]', style='Intense Quote')
            
            doc.add_paragraph()  # Spacing between sections
        
        # Add footer with page numbers
        footer_section = doc.sections[0]
        footer = footer_section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = 'Page '
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number field
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        footer_para.add_run(' | Confidential')
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_process_name = re.sub(r'[^\w\s-]', '', process_name).strip().replace(' ', '_')
        filename = f'Process_Analysis_{safe_process_name}_{timestamp}.docx'
        filepath = os.path.join('generated_reports', filename)
        
        # Save document
        doc.save(filepath)
        logger.info(f"Successfully generated Word document: {filename}")
        
        return filename
        
    except Exception as e:
        logger.error(f"Error generating Word document: {e}")
        raise


def list_generated_reports():
    """
    List all generated Word documents in the generated_reports folder.
    
    Returns:
        list: List of dictionaries with file information
    """
    try:
        reports_dir = 'generated_reports'
        if not os.path.exists(reports_dir):
            return []
        
        files = []
        for filename in os.listdir(reports_dir):
            if filename.endswith('.docx'):
                filepath = os.path.join(reports_dir, filename)
                stat = os.stat(filepath)
                files.append({
                    'filename': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                })
        
        # Sort by modified time, newest first
        files.sort(key=lambda x: x['modified'], reverse=True)
        return files
        
    except Exception as e:
        logger.error(f"Error listing generated reports: {e}")
        return []


def cleanup_old_reports(hours=24):
    """
    Delete reports older than specified hours.
    
    Args:
        hours (int): Delete files older than this many hours
        
    Returns:
        int: Number of files deleted
    """
    try:
        reports_dir = 'generated_reports'
        if not os.path.exists(reports_dir):
            return 0
        
        deleted_count = 0
        current_time = datetime.now().timestamp()
        cutoff_time = current_time - (hours * 3600)
        
        for filename in os.listdir(reports_dir):
            if filename.endswith('.docx'):
                filepath = os.path.join(reports_dir, filename)
                file_time = os.path.getmtime(filepath)
                
                if file_time < cutoff_time:
                    os.remove(filepath)
                    deleted_count += 1
                    logger.info(f"Deleted old report: {filename}")
        
        return deleted_count
        
    except Exception as e:
        logger.error(f"Error cleaning up old reports: {e}")
        return 0


# Backward compatibility alias
def create_sox_word_document(analysis_text, control_name='SOX Control Analysis', metadata=None):
    """
    Legacy function for backward compatibility.
    Calls create_process_document with renamed parameters.
    """
    return create_process_document(analysis_text, process_name=control_name, metadata=metadata)
